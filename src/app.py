import shutil
from pathlib import Path
from statistics import median
from typing import Any
from uuid import uuid4
import base64
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from src.agent import DocumentAgent
from src.layout_analyzer import LayoutAnalyzer
from src.ocr_engine import OCREngine
from src.reading_order import ordered_lines_from_structured, raw_text_from_structured, word_rows_from_structured
from src.settings import get_settings, save_settings
from src.table_extractor import TableExtractor

app = FastAPI()


class AnalyzeRequest(BaseModel):
    extracted_text: str = ""
    json_data: dict[str, Any] = Field(default_factory=dict)
    bounding_boxes: list[dict[str, Any]] = Field(default_factory=list)


class ChatRequest(BaseModel):
    question: str
    extracted_text: str = ""
    json_data: dict[str, Any] = Field(default_factory=dict)
    bounding_boxes: list[dict[str, Any]] = Field(default_factory=list)
    ai_analysis: dict[str, Any] | None = None


class LayoutRequest(BaseModel):
    json_data: dict[str, Any] = Field(default_factory=dict)
    tables: list[dict[str, Any]] = Field(default_factory=list)


class SettingsRequest(BaseModel):
    image_preprocessing_enabled: bool = True
    theme: str = "light"


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

engine = OCREngine(config_path="configs/model_config.yaml")
agent = DocumentAgent()
layout_analyzer = LayoutAnalyzer()
table_extractor = TableExtractor()


@app.get("/api/settings")
async def read_settings():
    return get_settings()


@app.put("/api/settings")
async def update_settings(payload: SettingsRequest):
    return save_settings(payload.dict())


@app.get("/api/layoutxlm/status")
async def layoutxlm_status():
    return agent.layoutxlm_extractor.status()


@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    original_filename = file.filename or "upload"
    extension = Path(original_filename).suffix.lower()
    allowed_extensions = {".jpg", ".jpeg", ".png", ".pdf"}
    if extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail="Chi ho tro file PDF, PNG, JPG, JPEG.",
        )

    upload_dir = Path("data/sample_images")
    upload_dir.mkdir(parents=True, exist_ok=True)
    temp_file_path = upload_dir / f"temp_{uuid4().hex}{extension}"

    try:
        with temp_file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        raw_result = engine.process_document(str(temp_file_path))

        processed_temp_path = getattr(engine, 'current_processed_image_path', None)
        base64_image = ""
        
        if processed_temp_path and os.path.exists(processed_temp_path):
            with open(processed_temp_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                base64_image = f"data:image/jpeg;base64,{encoded_string}"

    except Exception as exc:
        if temp_file_path.exists():
            temp_file_path.unlink()
        raise HTTPException(
            status_code=500,
            detail=f"Khong the xu ly OCR: {exc}",
        ) from exc

    finally:
        if temp_file_path.exists():
            temp_file_path.unlink()
            
        processed_temp_path_obj = Path(getattr(engine, 'current_processed_image_path', "")) if getattr(engine, 'current_processed_image_path', None) else None
        if processed_temp_path_obj and processed_temp_path_obj.exists() and processed_temp_path_obj != temp_file_path:
            try:
                processed_temp_path_obj.unlink()
            except OSError:
                pass

    structured_data = engine.get_structured_data(raw_result)
    structured_data["_processing"] = engine.get_processing_info()
    extracted_text = engine.get_raw_text(raw_result, structured_data=structured_data)
    tables = table_extractor.extract_tables(structured_data)
    layout_regions = layout_analyzer.analyze(structured_data, tables=tables)
    bounding_boxes = engine.get_bounding_boxes(structured_data)
    page_images = engine.take_document_images()

    ai_analysis = agent.analyze(
        extracted_text=extracted_text,
        structured_data=structured_data,
        bounding_boxes=bounding_boxes,
        page_images=page_images,
    )
    ai_analysis["layout_regions"] = layout_regions

    return {
        "status": "success",
        "filename": file.filename,
        "extracted_text": extracted_text,
        "json_data": structured_data,
        "tables": tables,
        "layout_regions": layout_regions,
        "ai_analysis": ai_analysis,
        "bounding_boxes": bounding_boxes,
        "processed_image_base64": base64_image
    }


@app.post("/api/analyze")
async def analyze_document(payload: AnalyzeRequest):
    analysis = agent.analyze(
        extracted_text=payload.extracted_text,
        structured_data=payload.json_data,
        bounding_boxes=payload.bounding_boxes,
    )
    analysis["layout_regions"] = layout_analyzer.analyze(payload.json_data)
    return analysis


@app.post("/api/chat")
async def chat_document(payload: ChatRequest):
    return agent.answer_question(
        question=payload.question,
        extracted_text=payload.extracted_text,
        analysis=payload.ai_analysis,
        structured_data=payload.json_data,
        bounding_boxes=payload.bounding_boxes,
    )


@app.post("/api/layout")
async def layout_document(payload: LayoutRequest):
    return {
        "layout_regions": layout_analyzer.analyze(
            payload.json_data,
            tables=payload.tables,
        )
    }


@app.post("/api/export-pdf")
async def export_pdf(data: dict):
    pdf_path = "data/exported_document.pdf"
    c = canvas.Canvas(pdf_path, pagesize=A4)
    page_width, page_height = A4
    font_name = _pdf_font_name()

    pages = data.get("pages", [])
    if not pages:
        _draw_wrapped_pdf_text(c, "No OCR text available.", page_width, page_height, font_name)
    else:
        special_handling = _uses_special_handling(data)
        for page_index, page in enumerate(pages):
            page_data = {"pages": [page]}
            if special_handling:
                rows = word_rows_from_structured(page_data).get(1, [])
                if rows:
                    _draw_corrected_rows_pdf_page(c, rows, page_width, page_height, font_name)
                else:
                    page_text = raw_text_from_structured(page_data)
                    if not page_text:
                        page_text = "\n".join(
                            line["text"]
                            for line in ordered_lines_from_structured(page_data)
                            if line.get("text")
                        )
                    _draw_wrapped_pdf_text(c, page_text, page_width, page_height, font_name)
            else:
                _draw_default_ocr_pdf_page(c, page_data, page_width, page_height, font_name)

            if page_index < len(pages) - 1:
                c.showPage()

    c.save()
    return FileResponse(pdf_path, media_type="application/pdf", filename="exported_document.pdf")


def _uses_special_handling(data: dict[str, Any]) -> bool:
    processing = data.get("_processing") if isinstance(data.get("_processing"), dict) else {}
    return bool(processing.get("special_handling"))


def _pdf_font_name() -> str:
    for font_name, font_path in (
        ("Arial", Path("C:/Windows/Fonts/arial.ttf")),
        ("DejaVuSans", Path("C:/Windows/Fonts/DejaVuSans.ttf")),
    ):
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            return font_name
        except Exception:
            continue
    return "Helvetica"


def _draw_wrapped_pdf_text(
    c: canvas.Canvas,
    text: str,
    page_width: float,
    page_height: float,
    font_name: str,
) -> None:
    margin = 42
    font_size = 10.5
    line_height = font_size * 1.35
    max_width = page_width - (margin * 2)
    y = page_height - margin

    c.setFont(font_name, font_size)
    for source_line in text.splitlines() or [""]:
        wrapped_lines = _wrap_pdf_line(c, source_line, font_name, font_size, max_width)
        if not wrapped_lines:
            wrapped_lines = [""]

        for line in wrapped_lines:
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = page_height - margin

            if line:
                c.drawString(margin, y, line)
            y -= line_height


def _draw_default_ocr_pdf_page(
    c: canvas.Canvas,
    page_data: dict[str, Any],
    page_width: float,
    page_height: float,
    font_name: str,
) -> None:
    lines = ordered_lines_from_structured(page_data)
    if not lines:
        page_text = raw_text_from_structured(page_data)
        _draw_wrapped_pdf_text(c, page_text or "No OCR text available.", page_width, page_height, font_name)
        return

    for line in lines:
        if not all(key in line for key in ("x1", "y1", "y2")):
            continue

        x = float(line["x1"]) * page_width
        y = page_height - (float(line["y1"]) * page_height)
        box_height = max(float(line["y2"]) - float(line["y1"]), 0.001) * page_height
        font_size = _clamp(box_height * 0.8, 6.0, 11.5)

        c.setFont(font_name, font_size)
        c.drawString(x, y - font_size, line["text"])


def _draw_corrected_rows_pdf_page(
    c: canvas.Canvas,
    rows: list[list[dict[str, Any]]],
    page_width: float,
    page_height: float,
    font_name: str,
) -> None:
    margin = 42
    font_size = 9.5
    line_height = font_size * 1.35
    max_width = page_width - (margin * 2)
    y = page_height - margin
    cleaned_rows = _clean_ocr_rows(rows)
    median_step = _median_row_step(cleaned_rows)
    previous_y1: float | None = None

    c.setFont(font_name, font_size)
    for row in cleaned_rows:
        row_y1, _ = _row_y_bounds(row)
        if previous_y1 is not None and _is_paragraph_gap(row_y1, previous_y1, median_step):
            y -= line_height * 0.85

        text = _docx_row_text(row)
        wrapped_lines = _wrap_pdf_line(c, text, font_name, font_size, max_width)
        for line in wrapped_lines:
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = page_height - margin

            c.drawString(margin, y, line)
            y -= line_height

        previous_y1 = row_y1


def _clamp(value: float, minimum: float, maximum: float) -> float:
    return max(minimum, min(value, maximum))


def _clean_ocr_rows(rows: list[list[dict[str, Any]]]) -> list[list[dict[str, Any]]]:
    cleaned = []
    for row in rows:
        row_words = sorted(
            [word for word in row if word.get("text")],
            key=lambda word: word.get("x1", 0.0),
        )
        if row_words:
            cleaned.append(row_words)
    return cleaned


def _row_y_bounds(row: list[dict[str, Any]]) -> tuple[float, float]:
    y1 = min(float(word.get("y1", 0.0)) for word in row)
    y2 = max(float(word.get("y2", y1)) for word in row)
    return y1, y2


def _median_row_step(rows: list[list[dict[str, Any]]]) -> float:
    row_tops = [_row_y_bounds(row)[0] for row in rows]
    deltas = [
        current - previous
        for previous, current in zip(row_tops, row_tops[1:])
        if current - previous > 0.002
    ]
    if not deltas:
        return 0.035
    return median(deltas)


def _is_paragraph_gap(current_y1: float, previous_y1: float, median_step: float) -> bool:
    return current_y1 - previous_y1 > max(median_step * 1.35, median_step + 0.012)


def _wrap_pdf_line(
    c: canvas.Canvas,
    text: str,
    font_name: str,
    font_size: float,
    max_width: float,
) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if c.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
            continue

        if current:
            lines.append(current)
            current = ""

        if c.stringWidth(word, font_name, font_size) <= max_width:
            current = word
        else:
            chunks = _split_pdf_word(c, word, font_name, font_size, max_width)
            lines.extend(chunks[:-1])
            current = chunks[-1] if chunks else ""

    if current:
        lines.append(current)
    return lines


def _split_pdf_word(
    c: canvas.Canvas,
    word: str,
    font_name: str,
    font_size: float,
    max_width: float,
) -> list[str]:
    chunks = []
    current = ""
    for char in word:
        candidate = current + char
        if c.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
            continue
        if current:
            chunks.append(current)
        current = char
    if current:
        chunks.append(current)
    return chunks


@app.post("/api/export-docx")
async def export_docx(data: dict):
    docx_path = "data/exported_document.docx"
    doc = Document()
    _configure_docx_page(doc)

    pages = data.get("pages", [])
    if not pages:
        _add_docx_plain_text(doc, "No OCR text available.")

    special_handling = _uses_special_handling(data)
    for page_index, page in enumerate(pages):
        page_data = {"pages": [page]}
        rows = word_rows_from_structured(page_data).get(1, []) if special_handling else []
        if special_handling and rows:
            _add_docx_corrected_rows(doc, rows)
        else:
            _add_docx_default_lines(doc, ordered_lines_from_structured(page_data))

        if page_index < len(pages) - 1:
            doc.add_page_break()

    doc.save(docx_path)
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="exported_document.docx",
    )


def _add_docx_default_lines(doc: Document, lines: list[dict[str, Any]]) -> None:
    for line in lines:
        if "x1" not in line:
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(line["x1"] * 6.5)
        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(line["text"])
        _set_docx_run_font(run, "Arial", 10)


def _add_docx_corrected_rows(doc: Document, rows: list[list[dict[str, Any]]]) -> None:
    font_name = "Arial"
    font_size = 10.5
    cleaned_rows = _clean_ocr_rows(rows)
    median_step = _median_row_step(cleaned_rows)
    previous_y1: float | None = None

    for row in cleaned_rows:
        row_y1, _ = _row_y_bounds(row)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if previous_y1 is not None and _is_paragraph_gap(row_y1, previous_y1, median_step):
            paragraph.paragraph_format.space_before = Pt(font_size * 0.85)
        else:
            paragraph.paragraph_format.space_before = Pt(0)

        run = paragraph.add_run(_docx_row_text(row))
        _set_docx_run_font(run, font_name, font_size)
        previous_y1 = row_y1


def _configure_docx_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def _add_docx_layout_rows(doc: Document, rows: list[list[dict[str, Any]]]) -> None:
    usable_width_inches = 8.27 - (0.55 * 2)
    usable_height_points = (11.69 - (0.55 * 2)) * 72
    font_name = "Arial"
    font_size = _estimate_docx_font_size(rows, usable_height_points)
    line_height = font_size * 1.25
    previous_y = 0.0

    for row_index, row in enumerate(rows):
        row = sorted(row, key=lambda word: word.get("x1", 0.0))
        text = _docx_row_text(row)
        if not text:
            continue

        x1 = min(float(word.get("x1", 0.0)) for word in row)
        y1 = min(float(word.get("y1", 0.0)) for word in row)
        if row_index == 0:
            space_before = min(max(y1 * usable_height_points, 0), 54)
        else:
            gap_points = max((y1 - previous_y) * usable_height_points - line_height, 0)
            space_before = min(gap_points, 30)

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(min(max(x1 * usable_width_inches, 0), usable_width_inches - 0.4))
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run(text)
        _set_docx_run_font(run, font_name, font_size)
        previous_y = max(float(word.get("y2", y1)) for word in row)


def _add_docx_plain_text(doc: Document, text: str) -> None:
    font_name = "Arial"
    for source_line in text.splitlines() or [""]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(source_line)
        _set_docx_run_font(run, font_name, 10.5)


def _docx_row_text(row: list[dict[str, Any]]) -> str:
    if not row:
        return ""

    widths = [float(word.get("width", 0.0)) for word in row if word.get("width")]
    median_width = median(widths) if widths else 0.015
    parts = [str(row[0].get("text", "")).strip()]

    for previous, current in zip(row, row[1:]):
        gap = float(current.get("x1", 0.0)) - float(previous.get("x2", 0.0))
        if gap > median_width * 2.2:
            spaces = min(max(int(gap / max(median_width, 0.001)), 2), 10)
            parts.append(" " * spaces)
        else:
            parts.append(" ")
        parts.append(str(current.get("text", "")).strip())

    return "".join(parts).strip()


def _estimate_docx_font_size(rows: list[list[dict[str, Any]]], usable_height_points: float) -> float:
    heights = [
        float(word.get("height", 0.0)) * usable_height_points
        for row in rows
        for word in row
        if word.get("height")
    ]
    if not heights:
        return 10.5
    return min(max(median(heights) * 0.78, 8.5), 12.5)


def _set_docx_run_font(run, font_name: str, font_size: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
